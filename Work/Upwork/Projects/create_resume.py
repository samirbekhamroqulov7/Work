from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# === HEADER ===
name = doc.add_paragraph()
name_run = name.add_run('SAMIRBEK')
name_run.bold = True
name_run.font.size = Pt(26)
name_run.font.color.rgb = RGBColor(30, 41, 59)
name.alignment = WD_ALIGN_PARAGRAPH.LEFT
name.space_after = Pt(2)

title = doc.add_paragraph()
title_run = title.add_run('Full-Stack Developer | Python, C#, PHP, JavaScript | Unity, WordPress, Automation')
title_run.font.size = Pt(12)
title_run.font.color.rgb = RGBColor(37, 99, 235)
title.space_after = Pt(4)

contacts = doc.add_paragraph()
contacts_run = contacts.add_run('Uzbekistan  |  English, Russian, Uzbek  |  Freelance (Upwork, Kwork)')
contacts_run.font.size = Pt(10)
contacts_run.font.color.rgb = RGBColor(100, 116, 139)
contacts.space_after = Pt(12)

# Divider
def add_divider():
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    run = p.add_run('_' * 80)
    run.font.color.rgb = RGBColor(226, 232, 240)
    run.font.size = Pt(8)

# Section header
def add_section_header(text):
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    p.space_after = Pt(8)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(37, 99, 235)

# === SUMMARY ===
add_section_header('Professional Summary')
summary = doc.add_paragraph()
summary.add_run('Versatile full-stack developer with experience in web applications, desktop software, mobile apps, game development, and process automation. Proficient in multiple programming languages and frameworks. Focused on delivering clean, functional solutions on time and within budget.')
summary.paragraph_format.space_after = Pt(6)

# === SKILLS ===
add_section_header('Technical Skills')

skills = [
    ('Languages', 'Python, C#, C++, PHP, JavaScript, TypeScript, HTML/CSS, SQL, VBA'),
    ('Web Development', 'WordPress, WooCommerce, Django, Flask, FastAPI, React, Node.js, REST API'),
    ('Desktop & Mobile', '.NET (WinForms, WPF, MAUI), Visual Studio, SQL Server, SQLite'),
    ('Game Development', 'Unity, C#, Game Mechanics, UI/UX, Physics, 2D/3D'),
    ('Automation & Bots', 'Telegram Bots (Aiogram, Telethon), Excel/VBA, Web Scraping, Selenium'),
    ('Databases', 'PostgreSQL, MySQL, SQL Server, SQLite, MongoDB'),
    ('Tools', 'Git, Docker, Linux, VS Code, Visual Studio, Figma'),
    ('Integrations', 'YooKassa, Yandex Maps API, 1C, Stripe, Payment Gateways'),
]

for label, items in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run_label = p.add_run(f'{label}: ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_items = p.add_run(items)
    run_items.font.size = Pt(10)
    run_items.font.color.rgb = RGBColor(71, 85, 105)

# === PROJECTS ===
add_section_header('Key Projects')

projects = [
    ('Unity Games (3 titles)', 'Unity, C#, Game Design',
     'Developed 3 complete games from concept to release. Implemented game mechanics, physics systems, UI, animations, and optimization for target platforms.'),
    ('Exam Database Application', 'C#, .NET WinForms, SQL Server',
     'Desktop and mobile application for exam preparation. Features: test database with categories, scoring system, progress tracking, timer, and results history.'),
    ('Multi-Vendor Marketplace', 'WordPress, WooCommerce, Dokan Pro, PHP, YooKassa API',
     'E-commerce marketplace with vendor registration, product catalog, payment processing, seller payouts, Yandex Maps integration, and 1C synchronization.'),
    ('Warehouse Management System', 'Excel, VBA, Macros',
     'Automated inventory tracking with multi-warehouse support, material flow tracking, production planning, deficit calculation, and operations log.'),
    ('Telegram Message Analysis System', 'Python, PostgreSQL, pgvector, FastAPI, NLP',
     'Semantic search system across Telegram chat history. Uses embeddings for relevance scoring, finds potential leads based on service topic input.'),
]

for name, tech, desc in projects:
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(8)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(11)
    
    p_tech = doc.add_paragraph()
    p_tech.paragraph_format.space_after = Pt(2)
    run_tech = p_tech.add_run(tech)
    run_tech.font.size = Pt(10)
    run_tech.font.color.rgb = RGBColor(37, 99, 235)
    run_tech.italic = True
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(4)
    run_desc = p_desc.add_run(desc)
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = RGBColor(71, 85, 105)

# === EXPERIENCE ===
add_section_header('Experience')

p_exp = doc.add_paragraph()
p_exp.paragraph_format.space_after = Pt(2)
run_title = p_exp.add_run('Freelance Full-Stack Developer')
run_title.bold = True
run_title.font.size = Pt(11)
run_date = p_exp.add_run('    2023 — Present')
run_date.font.size = Pt(10)
run_date.font.color.rgb = RGBColor(148, 163, 184)

bullets = [
    'Web development: WordPress sites, custom PHP/Python backends, REST APIs',
    'Desktop applications: .NET WinForms, database-driven tools',
    'Game development: Unity projects with C#',
    'Automation: Telegram bots, Excel/VBA systems, web scrapers',
    'Client communication, requirements gathering, on-time delivery',
]

for bullet in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run(bullet)
    if not p.runs:
        pass
    else:
        run.font.size = Pt(10)
    # Clear and re-add
    p.clear()
    run = p.add_run(bullet)
    run.font.size = Pt(10)

# === LANGUAGES ===
add_section_header('Languages')

langs = doc.add_paragraph()
langs.add_run('Uzbek').bold = True
langs.add_run(' — Native    |    ')
langs.add_run('Russian').bold = True
langs.add_run(' — Fluent    |    ')
langs.add_run('English').bold = True
langs.add_run(' — Conversational')

# === EDUCATION ===
add_section_header('Education')
edu = doc.add_paragraph()
edu.add_run('[Your University / College]').bold = True
edu.add_run(' — Computer Science / IT')
edu2 = doc.add_paragraph()
run_edu2 = edu2.add_run('[Year — Year]')
run_edu2.font.color.rgb = RGBColor(148, 163, 184)
run_edu2.font.size = Pt(10)

# Save
filepath = r'c:\Work\Upwork\Projects\Samirbek_Resume.docx'
doc.save(filepath)
print(f'Resume saved: {filepath}')
